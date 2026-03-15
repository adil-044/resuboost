import fitz  # PyMuPDF
import docx
from io import BytesIO

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(stream=content, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX including headers and footers."""
    doc = docx.Document(BytesIO(content))
    full_text = []
    
    # 1. Extract from Headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
    # 2. Extract from Main Body
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # 3. Extract from Footers
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
    return "\n".join(full_text)

def parse_file(content: bytes, filename: str) -> str:
    """Detect file type and extract text."""
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(content)
    else:
        return ""
